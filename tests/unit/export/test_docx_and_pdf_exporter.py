"""Tests for DOCX and PDF exporters."""

from pathlib import Path

import pytest


pytest.importorskip("docx")

from docx import Document  # noqa: E402

from bloginator.export.docx_exporter import DOCXExporter  # noqa: E402
from bloginator.export.pdf_exporter import _REPORTLAB_AVAILABLE, PDFExporter  # noqa: E402
from bloginator.models.draft import Citation, Draft, DraftSection  # noqa: E402
from bloginator.models.outline import Outline, OutlineSection  # noqa: E402


REPORTLAB_AVAILABLE = _REPORTLAB_AVAILABLE


def _create_nested_draft() -> Draft:
    child = DraftSection(
        title="Child Section",
        content="Child content",
        citations=[],
    )
    parent = DraftSection(
        title="Parent Section",
        content="Parent content",
        citations=[
            Citation(
                chunk_id="c1",
                document_id="d1",
                filename="doc1.md",
                content_preview="Preview",
                similarity_score=0.9,
            )
        ],
        subsections=[child],
    )
    draft = Draft(
        title="Draft Title",
        thesis="Draft thesis",
        keywords=["testing"],
        sections=[parent],
    )
    draft.calculate_stats()
    return draft


def _create_nested_outline() -> Outline:
    child = OutlineSection(
        title="Child Section",
        description="Child description",
        coverage_pct=60.0,
        source_count=1,
        notes="Child note",
    )
    parent = OutlineSection(
        title="Parent Section",
        description="Parent description",
        coverage_pct=80.0,
        source_count=2,
        notes="Parent note",
        subsections=[child],
    )
    outline = Outline(
        title="Outline Title",
        thesis="Outline thesis",
        keywords=["testing"],
        sections=[parent],
    )
    outline.calculate_stats()
    return outline


def test_docx_exporter_creates_document_with_content(tmp_path: Path) -> None:
    draft = _create_nested_draft()
    exporter = DOCXExporter()
    output_path = tmp_path / "draft.docx"

    exporter.export_draft(draft, output_path)

    assert output_path.exists()
    doc = Document(output_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Draft Title" in text
    assert "Parent Section" in text
    assert "Child Section" in text
    assert "[1 sources]" in text


def test_docx_exporter_writes_outline_with_coverage_and_notes(tmp_path: Path) -> None:
    outline = _create_nested_outline()
    exporter = DOCXExporter()
    output_path = tmp_path / "outline.docx"

    exporter.export_outline(outline, output_path)

    assert output_path.exists()
    doc = Document(output_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Parent Section" in text
    assert "Child Section" in text
    assert "Coverage: 80%" in text
    assert "2 document(s)" in text
    assert "Parent note" in text
    assert "Child note" in text


@pytest.mark.skipif(not REPORTLAB_AVAILABLE, reason="requires reportlab")
def test_pdf_exporter_generates_non_empty_files(tmp_path: Path) -> None:
    draft = _create_nested_draft()
    outline = _create_nested_outline()
    exporter = PDFExporter()

    draft_path = tmp_path / "draft.pdf"
    outline_path = tmp_path / "outline.pdf"

    exporter.export_draft(draft, draft_path)
    exporter.export_outline(outline, outline_path)

    assert draft_path.exists() and draft_path.stat().st_size > 0
    assert outline_path.exists() and outline_path.stat().st_size > 0


@pytest.mark.skipif(REPORTLAB_AVAILABLE, reason="requires environment without reportlab")
def test_pdf_exporter_raises_runtime_error_when_reportlab_missing(tmp_path: Path) -> None:
    draft = _create_nested_draft()
    exporter = PDFExporter()

    with pytest.raises(RuntimeError):
        exporter.export_draft(draft, tmp_path / "draft.pdf")

"""DOCX exporter for drafts and outlines using python-docx."""

from __future__ import annotations

from typing import TYPE_CHECKING, Any


try:
    from docx import Document as _DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - optional dependency path
    _DOCX_AVAILABLE = False

from bloginator.export.base import Exporter


if TYPE_CHECKING:  # pragma: no cover - imports used only for type checking
    from pathlib import Path

    from bloginator.models.draft import Draft, DraftSection
    from bloginator.models.outline import Outline, OutlineSection


class DOCXExporter(Exporter):
    """Exports documents to Microsoft Word DOCX format."""

    def export_draft(self, draft: Draft, output_path: Path) -> None:
        """Export draft to DOCX file.

        Args:
            draft: Draft document to export
            output_path: Path where DOCX file should be saved
        """
        if not _DOCX_AVAILABLE:
            raise RuntimeError(
                "DOCX export requires the 'python-docx' package. "
                "Install bloginator[export] to enable this feature."
            )

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Create document
        doc = _DocxDocument()
        self._setup_styles(doc)

        # Title
        title = doc.add_heading(draft.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Thesis
        if draft.thesis:
            thesis_para = doc.add_paragraph()
            thesis_run = thesis_para.add_run(draft.thesis)
            thesis_run.italic = True
            thesis_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Spacing

        # Metadata
        metadata = doc.add_paragraph()
        metadata_text = (
            f"Classification: {draft.classification} | Audience: {draft.audience} | "
            f"Generated: {draft.created_date.strftime('%Y-%m-%d')}"
        )
        meta_run = metadata.add_run(metadata_text)
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Horizontal rule
        doc.add_paragraph("_" * 80)
        doc.add_paragraph()  # Spacing

        # Sections
        for section in draft.sections:
            self._add_draft_section(doc, section, level=1)

        # Save
        doc.save(str(output_path))

    def export_outline(self, outline: Outline, output_path: Path) -> None:
        """Export outline to DOCX file.

        Args:
            outline: Outline document to export
            output_path: Path where DOCX file should be saved
        """
        if not _DOCX_AVAILABLE:
            raise RuntimeError(
                "DOCX export requires the 'python-docx' package. "
                "Install bloginator[export] to enable this feature."
            )

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Create document
        doc = _DocxDocument()
        self._setup_styles(doc)

        # Title
        title = doc.add_heading(outline.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Thesis
        if outline.thesis:
            thesis_para = doc.add_paragraph()
            thesis_run = thesis_para.add_run(outline.thesis)
            thesis_run.italic = True
            thesis_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Spacing

        # Metadata
        metadata = doc.add_paragraph()
        metadata_text = (
            f"Classification: {outline.classification} | Audience: {outline.audience} | "
            f"Created: {outline.created_date.strftime('%Y-%m-%d')}"
        )
        meta_run = metadata.add_run(metadata_text)
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Horizontal rule
        doc.add_paragraph("_" * 80)
        doc.add_paragraph()  # Spacing

        # Sections
        for section in outline.sections:
            self._add_outline_section(doc, section, level=1)

        # Save
        doc.save(str(output_path))

    def get_file_extension(self) -> str:
        """Get file extension for DOCX format.

        Returns:
            'docx'
        """
        return "docx"

    def _setup_styles(self, doc: Any) -> None:
        """Setup custom styles for the document.

        Args:
            doc: Document-like object to add styles to
        """
        styles = doc.styles

        # Heading 1 style (blue color)
        try:
            heading1 = styles["Heading 1"]
            heading1.font.color.rgb = RGBColor(30, 136, 229)  # Blue
        except KeyError:
            # If the style is missing we fall back to the library defaults.
            pass

    def _add_draft_section(self, doc: Any, section: DraftSection, level: int) -> None:
        """Add draft section to document.

        Args:
            doc: Document-like object to add section to
            section: Draft section to add
            level: Heading level (1-9)
        """
        # Section heading
        doc.add_heading(section.title, level=min(level, 9))

        # Content
        if section.content:
            doc.add_paragraph(section.content)

            # Citation note
            if section.citations:
                citation_para = doc.add_paragraph()
                citation_run = citation_para.add_run(f"[{len(section.citations)} sources]")
                citation_run.italic = True
                citation_run.font.size = Pt(9)
                citation_run.font.color.rgb = RGBColor(128, 128, 128)

        # Subsections
        for subsection in section.subsections:
            self._add_draft_section(doc, subsection, level + 1)

    def _add_outline_section(self, doc: Any, section: OutlineSection, level: int) -> None:
        """Add outline section to document.

        Args:
            doc: Document-like object to add section to
            section: Outline section to add
            level: Heading level (1-9)
        """
        # Section heading
        doc.add_heading(section.title, level=min(level, 9))

        # Description
        if section.description:
            doc.add_paragraph(section.description)

        # Coverage and sources
        coverage_text = (
            f"Coverage: {section.coverage_pct:.0f}% from {section.source_count} document(s)"
        )
        coverage_para = doc.add_paragraph(coverage_text)
        if coverage_para.runs:
            coverage_run = coverage_para.runs[0]
            coverage_run.font.size = Pt(9)
            coverage_run.font.color.rgb = RGBColor(128, 128, 128)

        # Notes
        if section.notes:
            notes_para = doc.add_paragraph(f"Note: {section.notes}")
            if notes_para.runs:
                notes_run = notes_para.runs[0]
                notes_run.italic = True
                notes_run.font.size = Pt(9)
                notes_run.font.color.rgb = RGBColor(128, 128, 128)

        # Subsections
        for subsection in section.subsections:
            self._add_outline_section(doc, subsection, level + 1)

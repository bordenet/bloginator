"""Document content extractors for various file formats."""

import os
import re
import shutil
import subprocess
import sys
import tempfile
from collections.abc import Generator
from contextlib import contextmanager
from pathlib import Path

import fitz  # PyMuPDF


@contextmanager
def suppress_stderr() -> Generator[None, None, None]:
    """Context manager to suppress stderr output.

    This is useful for suppressing C-level warnings from PyMuPDF
    that clutter the output but are informational, not errors.
    """
    # Save original stderr
    original_stderr = sys.stderr
    original_stderr_fd = os.dup(2)

    try:
        # Redirect stderr to devnull
        devnull = os.open(os.devnull, os.O_WRONLY)
        os.dup2(devnull, 2)
        sys.stderr = os.fdopen(devnull, "w")

        yield

    finally:
        # Restore original stderr
        os.dup2(original_stderr_fd, 2)
        sys.stderr = original_stderr
        os.close(original_stderr_fd)


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text content from PDF file.

    Uses PyMuPDF (fitz) to extract text from all pages.

    Args:
        pdf_path: Path to PDF file

    Returns:
        Extracted text content

    Raises:
        FileNotFoundError: If PDF file does not exist
        ValueError: If file cannot be opened as PDF
    """
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    try:
        # Suppress MuPDF stderr warnings (syntax errors in PDF content streams)
        # These are informational C-level warnings, not extraction errors
        with suppress_stderr():
            doc = fitz.open(str(pdf_path))
            text_parts = []

            for page in doc:
                text_parts.append(page.get_text())

            doc.close()

        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}") from e


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text content from DOCX file.

    Uses python-docx to extract text from paragraphs and tables.

    Args:
        docx_path: Path to DOCX file

    Returns:
        Extracted text content

    Raises:
        FileNotFoundError: If DOCX file does not exist
        ValueError: If file cannot be opened as DOCX
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise ValueError("python-docx not installed") from e

    try:
        doc = DocxDocument(str(docx_path))
        text_parts = []

        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}") from e


def extract_text_from_doc(doc_path: Path) -> str:
    """Extract text content from legacy .doc file.

    Handles multiple .doc formats:
    1. MIME-encoded HTML (common from Confluence exports)
    2. Plain text files with .doc extension
    3. Real MS Word binary documents (via LibreOffice)
    4. Fallback to textract if available

    Args:
        doc_path: Path to legacy .doc file

    Returns:
        Extracted text content

    Raises:
        FileNotFoundError: If .doc file does not exist
        ValueError: If file cannot be converted or no converter available
    """
    if not doc_path.exists():
        raise FileNotFoundError(f".doc file not found: {doc_path}")

    try:
        with doc_path.open("rb") as f:
            # Read first few bytes to check file signature
            header = f.read(512)

            # Check for MS Word OLE signature (D0 CF 11 E0 A1 B1 1A E1)
            ole_signature = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

            if header.startswith(ole_signature):
                # This is a real MS Word document - use LibreOffice
                return _extract_real_doc_file(doc_path)

            # Check for MIME-encoded Confluence export
            f.seek(0)
            content = f.read()
            text = content.decode("utf-8", errors="replace")

            if "Exported From Confluence" in text or "multipart/related" in text:
                # This is a MIME-encoded Confluence export
                return _extract_confluence_export(text)

            # Otherwise, treat as plain text
            # Clean up any obvious non-content
            if text and any(c.isalpha() for c in text[:1000]):
                return text

            return text

    except Exception as e:
        raise ValueError(f"Failed to extract .doc file: {e}") from e


def _extract_confluence_export(mime_content: str) -> str:
    """Extract text from Confluence MIME-encoded export.

    Confluence exports pages as .doc files that are actually MIME-encoded
    HTML with quoted-printable encoding.

    Args:
        mime_content: Raw MIME content

    Returns:
        Extracted plain text
    """
    import quopri

    # Find the HTML content part
    html_content = ""

    # Look for quoted-printable HTML section
    parts = mime_content.split("------=_Part_")
    for part in parts:
        if "text/html" in part and "quoted-printable" in part.lower():
            # Find where the HTML content starts (after headers)
            lines = part.split("\n")
            content_start = 0
            for i, line in enumerate(lines):
                if line.strip() == "":
                    content_start = i + 1
                    break

            # Decode quoted-printable content
            encoded_content = "\n".join(lines[content_start:])
            try:
                html_bytes = quopri.decodestring(encoded_content.encode("utf-8"))
                html_content = html_bytes.decode("utf-8", errors="replace")
            except Exception:
                html_content = encoded_content
            break

    if not html_content:
        # Fallback: just strip everything that looks like email headers
        html_content = mime_content

    # Parse HTML to extract text
    return _html_to_text(html_content)


def _html_to_text(html_content: str) -> str:
    """Convert HTML to plain text.

    Args:
        html_content: HTML string

    Returns:
        Plain text with HTML tags removed
    """
    import html

    # Remove script and style elements
    text = re.sub(r"<script[^>]*>.*?</script>", "", html_content, flags=re.DOTALL | re.I)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.I)

    # Replace common block elements with newlines
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"<p[^>]*>", "\n\n", text, flags=re.I)
    text = re.sub(r"</p>", "", text, flags=re.I)
    text = re.sub(r"<div[^>]*>", "\n", text, flags=re.I)
    text = re.sub(r"</div>", "", text, flags=re.I)
    text = re.sub(r"<h[1-6][^>]*>", "\n\n", text, flags=re.I)
    text = re.sub(r"</h[1-6]>", "\n", text, flags=re.I)
    text = re.sub(r"<li[^>]*>", "\n• ", text, flags=re.I)
    text = re.sub(r"</li>", "", text, flags=re.I)
    text = re.sub(r"<tr[^>]*>", "\n", text, flags=re.I)
    text = re.sub(r"<td[^>]*>", " | ", text, flags=re.I)

    # Remove all remaining HTML tags
    text = re.sub(r"<[^>]+>", "", text)

    # Decode HTML entities
    text = html.unescape(text)

    # Clean up whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    return text


def _extract_real_doc_file(doc_path: Path) -> str:
    """Extract text from a real MS Word binary document.

    Args:
        doc_path: Path to MS Word .doc file

    Returns:
        Extracted text content

    Raises:
        ValueError: If conversion fails
    """
    # Try LibreOffice (soffice) first - most reliable
    soffice_path = shutil.which("soffice")
    if soffice_path:
        try:
            return _extract_doc_with_libreoffice(doc_path)
        except Exception:
            pass  # Fall through to try other methods

    # Try textract if available
    try:
        import textract

        result = textract.process(str(doc_path)).decode("utf-8")
        return str(result)
    except ImportError:
        pass
    except Exception as e:
        raise ValueError(f"textract failed to extract .doc: {e}") from e

    raise ValueError(
        f"Cannot extract binary .doc file: {doc_path}. "
        "Install LibreOffice (brew install --cask libreoffice) or textract."
    )


def _extract_doc_with_libreoffice(doc_path: Path) -> str:
    """Convert .doc to text using LibreOffice headless mode.

    Args:
        doc_path: Path to .doc file

    Returns:
        Extracted text content

    Raises:
        ValueError: If conversion fails
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        # Convert to text using LibreOffice
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                tmpdir,
                str(doc_path.absolute()),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise ValueError(f"LibreOffice conversion failed: {result.stderr}")

        # Find the output file
        txt_files = list(Path(tmpdir).glob("*.txt"))
        if not txt_files:
            raise ValueError("LibreOffice produced no output file")

        # Read the converted text
        text = txt_files[0].read_text(encoding="utf-8", errors="replace")
        return text


def extract_text_from_markdown(md_path: Path) -> str:
    """Extract text content from Markdown file.

    Reads file and optionally strips YAML frontmatter.

    Args:
        md_path: Path to Markdown file

    Returns:
        Extracted text content (with frontmatter removed)

    Raises:
        FileNotFoundError: If Markdown file does not exist
    """
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {md_path}")

    content = md_path.read_text(encoding="utf-8")

    # Remove YAML frontmatter
    frontmatter_pattern = r"^---\s*\n.*?\n---\s*\n"
    content = re.sub(frontmatter_pattern, "", content, flags=re.DOTALL)

    return content.strip()


def extract_text_from_txt(txt_path: Path) -> str:
    """Extract text content from plain text file.

    Args:
        txt_path: Path to text file

    Returns:
        Extracted text content

    Raises:
        FileNotFoundError: If text file does not exist
    """
    if not txt_path.exists():
        raise FileNotFoundError(f"Text file not found: {txt_path}")

    return txt_path.read_text(encoding="utf-8")


def extract_text_from_file(file_path: Path) -> str:
    """Extract text from file based on extension.

    Automatically detects file type and uses appropriate extractor.

    Args:
        file_path: Path to file

    Returns:
        Extracted text content

    Raises:
        FileNotFoundError: If file does not exist
        ValueError: If file type is not supported
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix == ".doc":
        return extract_text_from_doc(file_path)
    elif suffix in [".md", ".markdown"]:
        return extract_text_from_markdown(file_path)
    elif suffix in [".txt", ".text"]:
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def extract_section_headings(text: str) -> list[tuple[str, int]]:
    """Extract section headings from Markdown text.

    Finds Markdown headings (# Heading, ## Subheading, etc.) and
    their positions in the text.

    Args:
        text: Markdown text content

    Returns:
        List of (heading_text, char_position) tuples
    """
    headings = []
    pattern = r"^(#{1,6})\s+(.+)$"

    for match in re.finditer(pattern, text, re.MULTILINE):
        heading_text = match.group(2).strip()
        char_position = match.start()
        headings.append((heading_text, char_position))

    return headings
